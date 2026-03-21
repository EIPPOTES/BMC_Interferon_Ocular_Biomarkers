# 路径配置
from pathlib import Path
import sys
sys.path.insert(0, str(Path(__file__).parent.parent))
from configs.paths_config import *
import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os


def main():
    """主函数，包装原有执行代码"""
    print("=" * 80)
    print("将Markdown转换为Word格式")
    print("=" * 80)

    # 读取Markdown文件
    input_file = '/mnt/c/Users/CUI/Desktop/论文及图表/完整论文_含参考文献.md'
    output_file = '/mnt/c/Users/CUI/Desktop/论文及图表/完整论文_含参考文献.docx'

    with open(input_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    print(f"已读取Markdown文件: {len(md_content)} 字符")

    # 创建Word文档
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # 处理Markdown内容
    lines = md_content.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # 跳过空行
        if not line:
            i += 1
            continue

        # 标题处理
        if line.startswith('# Title'):
            i += 1
            continue
        elif line.startswith('# '):
            # 一级标题 - 论文标题
            title_text = line[2:].strip()
            if title_text and not title_text.startswith('Title'):
                heading = doc.add_heading(title_text, level=0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in heading.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(16)
                    run.font.bold = True

        elif line.startswith('## '):
            # 二级标题
            heading_text = line[3:].strip()
            if heading_text:
                heading = doc.add_heading(heading_text, level=1)
                for run in heading.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.font.bold = True

        elif line.startswith('### '):
            # 三级标题
            heading_text = line[4:].strip()
            if heading_text:
                heading = doc.add_heading(heading_text, level=2)
                for run in heading.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)
                    run.font.bold = True

        elif line.startswith('#### '):
            # 四级标题
            heading_text = line[5:].strip()
            if heading_text:
                heading = doc.add_heading(heading_text, level=3)
                for run in heading.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.bold = True

        elif line.startswith('---'):
            # 分隔线，跳过
            pass

        elif line.startswith('- [x]') or line.startswith('- [ ]'):
            # 复选框列表
            check_text = line[5:].strip()
            if check_text:
                p = doc.add_paragraph(check_text, style='List Bullet')
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

        elif line.startswith('- ') or line.startswith('* '):
            # 无序列表
            list_text = line[2:].strip()
            if list_text:
                p = doc.add_paragraph(list_text, style='List Bullet')
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

        elif re.match(r'^\d+\.', line):
            # 有序列表（参考文献）
            ref_text = line.strip()
            if ref_text:
                p = doc.add_paragraph(ref_text, style='List Number')
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

        elif line.startswith('|'):
            # 表格 - 跳过（提示用户手动插入）
            if '---' not in line and '特征' not in line and '指标' not in line:
                p = doc.add_paragraph(f"[Table: {line.strip('| ')}]")
                p.runs[0].font.italic = True
                p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        elif line.startswith('**') and line.endswith('**'):
            # 粗体文本
            bold_text = line.strip('*')
            p = doc.add_paragraph()
            run = p.add_run(bold_text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        else:
            # 普通段落
            # 处理行内格式
            p = doc.add_paragraph()

            # 处理粗体 **text**
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                else:
                    run = p.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

            # 设置段落格式
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)

        i += 1

    # 保存文档
    doc.save(output_file)

    print(f"\n✓ Word文档已生成!")
    print(f"  位置: {output_file}")
    print(f"  大小: {os.path.getsize(output_file) / 1024:.1f} KB")

    print("\n" + "=" * 80)
    print("转换说明:")
    print("=" * 80)
    print("""
    1. 标题已正确格式化（居中、加粗）
    2. 段落已设置行距和字体
    3. 列表已转换为Word列表格式
    4. 表格位置已标记为[Table]，请手动插入图片
    5. 参考文献已编号

    注意:
    - 请检查文档格式，可能需要微调
    - 表格需要手动插入为图片
    - 图表需要手动插入
    - 建议最终使用Word的"另存为PDF"功能生成PDF版本
    """)

    print("\n转换完成!")



if __name__ == "__main__":
    main()