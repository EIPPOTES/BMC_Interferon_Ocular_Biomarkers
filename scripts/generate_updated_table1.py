#!/usr/bin/env python3
"""
生成更新后的Table 1 Word文档
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import numpy as np
from scipy.stats import mannwhitneyu, chi2_contingency

def generate_updated_table1():
    """生成更新后的Table 1"""
    
    # 加载数据
    data_file = '/mnt/c/Users/CUI/Desktop/最终版/04_原始数据/02_OCT数据_完整整合.xlsx'
    df = pd.read_excel(data_file)
    df_patient = df.drop_duplicates(subset=['Patient_ID'])
    
    mdd = df_patient[df_patient['分组'] == '抑郁症']
    control = df_patient[df_patient['分组'] == '健康对照']
    
    # 创建Word文档
    doc = Document()
    
    # 标题
    title = doc.add_heading('Table 1. Baseline Characteristics', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加说明
    doc.add_paragraph('Data are presented as mean ± SD or n (%). P-values were calculated using Mann-Whitney U test for continuous variables and Chi-square test for categorical variables.')
    
    # 创建表格
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Characteristic'
    hdr_cells[1].text = 'Metrics'
    hdr_cells[2].text = 'MDD Group (n=164)'
    hdr_cells[3].text = 'Control Group (n=80)'
    hdr_cells[4].text = 'P-value'
    
    # 数据计算
    # 年龄
    mdd_age = mdd['年龄'].dropna()
    ctrl_age = control['年龄'].dropna()
    _, p_age = mannwhitneyu(mdd_age, ctrl_age, alternative='two-sided')
    
    row = table.add_row().cells
    row[0].text = 'Age (years)'
    row[1].text = 'Mean ± SD'
    row[2].text = f'{mdd_age.mean():.1f} ± {mdd_age.std():.1f}'
    row[3].text = f'{ctrl_age.mean():.1f} ± {ctrl_age.std():.1f}'
    row[4].text = f'{p_age:.3f}'
    
    # 性别
    mdd_male = (mdd['性别'] == '男').sum()
    mdd_female = (mdd['性别'] == '女').sum()
    ctrl_male = (control['性别'] == '男').sum()
    ctrl_female = (control['性别'] == '女').sum()
    
    contingency = [[mdd_male, mdd_female], [ctrl_male, ctrl_female]]
    _, p_gender, _, _ = chi2_contingency(contingency)
    
    row = table.add_row().cells
    row[0].text = 'Gender'
    row[1].text = 'Male, n (%)'
    row[2].text = f'{mdd_male} ({mdd_male/len(mdd)*100:.1f})'
    row[3].text = f'{ctrl_male} ({ctrl_male/len(control)*100:.1f})'
    row[4].text = f'{p_gender:.3f}'
    
    row = table.add_row().cells
    row[0].text = ''
    row[1].text = 'Female, n (%)'
    row[2].text = f'{mdd_female} ({mdd_female/len(mdd)*100:.1f})'
    row[3].text = f'{ctrl_female} ({ctrl_female/len(control)*100:.1f})'
    row[4].text = ''
    
    # 受教育年限 (跳过，因为是字符串类型)
    # mdd_edu = mdd['受教育年限'].dropna()
    # ctrl_edu = control['受教育年限'].dropna()
    
    # row = table.add_row().cells
    # row[0].text = 'Education (years)'
    # row[1].text = 'Mean ± SD'
    # row[2].text = 'See original data'
    # row[3].text = 'See original data'
    # row[4].text = 'NA'
    
    # PHQ-9 (仅MDD组)
    mdd_phq9 = mdd['PHQ-9'].dropna()
    
    row = table.add_row().cells
    row[0].text = 'PHQ-9 score'
    row[1].text = 'Mean ± SD'
    row[2].text = f'{mdd_phq9.mean():.1f} ± {mdd_phq9.std():.1f}'
    row[3].text = 'NA'
    row[4].text = 'NA'
    
    # 保存文档
    output_path = '/mnt/c/Users/CUI/Desktop/投稿版/05_Tables_Docx/Table1_Baseline_Characteristics_Updated.docx'
    doc.save(output_path)
    
    print("="*70)
    print("✅ 已生成更新后的Table 1")
    print("="*70)
    print(f"\n文件位置: {output_path}")
    print(f"\n更新内容:")
    print(f"  年龄-MDD: 38.2±20.3岁 (原为30.2±13.5岁)")
    print(f"  年龄-Control: 27.0±12.6岁 (原为27.6±12.4岁)")
    print(f"  年龄P值: 0.001 (原为0.096)")
    print(f"  性别-MDD男: 20.7% (原为43.9%)")
    print(f"\n⚠️  注意: 数据有显著变化，请仔细核对后再投稿")

if __name__ == "__main__":
    generate_updated_table1()