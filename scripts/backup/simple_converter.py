#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def convert_md_to_docx(md_path, docx_path):
    """Convert markdown to docx with basic formatting."""
    
    # Read markdown file
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
            
        # Handle title
        if line.startswith('# Title'):
            # Find the actual title
            title_found = False
            for j in range(i + 1, min(i + 10, len(lines))):
                if lines[j].strip() and not lines[j].strip().startswith('---'):
                    title_text = lines[j].strip()
                    title = doc.add_heading(title_text, level=0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in title.runs:
                        run.font.size = Pt(14)
                        run.font.bold = True
                    title_found = True
                    i = j + 1
                    break
            if not title_found:
                i += 1
            continue
                
        # Handle sections
        if line.startswith('# ') and not line.startswith('# Title'):
            section_title = line[2:].strip()
            if section_title == 'Abstract':
                doc.add_page_break()
            doc.add_heading(section_title, level=1)
            i += 1
            continue
            
        # Handle subsections
        if line.startswith('## '):
            subsection_title = line[3:].strip()
            doc.add_heading(subsection_title, level=2)
            i += 1
            continue
            
        # Handle sub-subsections
        if line.startswith('### '):
            subsubsection_title = line[4:].strip()
            doc.add_heading(subsubsection_title, level=3)
            i += 1
            continue
            
        # Handle table-like lines (simplified)
        if '|' in line and line.count('|') > 2:
            # Skip table separator lines
            if re.search(r'[-:]+', line.replace('|', '')):
                i += 1
                continue
            # Add table line as regular text
            p = doc.add_paragraph(line.replace('|', ' | '))
            i += 1
            continue
            
        # Regular paragraph
        # Check if this is part of a multi-line paragraph
        paragraph_text = line
        j = i + 1
        while j < len(lines) and lines[j].strip() and not lines[j].startswith('#') and not lines[j].startswith('##') and not lines[j].startswith('###'):
            paragraph_text += ' ' + lines[j].rstrip()
            j += 1
            
        if paragraph_text:
            p = doc.add_paragraph(paragraph_text)
            
        i = j
        
    # Save document
    doc.save(docx_path)
    print(f"Document saved to {docx_path}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) != 3:
        print("Usage: python simple_converter.py <input.md> <output.docx>")
        sys.exit(1)
    
    convert_md_to_docx(sys.argv[1], sys.argv[2])